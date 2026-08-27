import io

import pytest

from app import create_app


@pytest.fixture
def client(tmp_path):
    app = create_app(
        {
            "TESTING": True,
            "UPLOAD_FOLDER": str(tmp_path),
            "MAX_CONTENT_LENGTH": 5 * 1024 * 1024,
        }
    )

    return app.test_client()


def test_upload_without_file(client):
    response = client.post("/api/resume/upload")

    assert response.status_code == 400
    assert response.get_json()["error"] == "No resume file provided"


def test_upload_without_filename(client):
    data = {
        "resume": (io.BytesIO(b"test content"), ""),
    }

    response = client.post(
        "/api/resume/upload",
        data=data,
        content_type="multipart/form-data",
    )

    assert response.status_code == 400
    assert response.get_json()["error"] == "No file selected"


def test_upload_unsupported_file(client):
    data = {
        "resume": (io.BytesIO(b"plain text"), "resume.txt"),
    }

    response = client.post(
        "/api/resume/upload",
        data=data,
        content_type="multipart/form-data",
    )

    assert response.status_code == 400
    assert "PDF and DOCX" in response.get_json()["error"]


def test_upload_valid_docx(client):
    from docx import Document

    document = Document()
    document.add_heading("John Doe", level=1)
    document.add_heading("SKILLS", level=2)
    document.add_paragraph("Python")
    document.add_paragraph("SQL")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    data = {
        "resume": (
            buffer,
            "test_resume.docx",
        ),
    }

    response = client.post(
        "/api/resume/upload",
        data=data,
        content_type="multipart/form-data",
    )

    assert response.status_code == 200

    result = response.get_json()

    assert result["message"] == "Resume processed successfully"
    assert result["filename"] == "test_resume.docx"
    assert "resume" in result
    assert "sections" in result["resume"]
    assert "skills" in result["resume"]["sections"]